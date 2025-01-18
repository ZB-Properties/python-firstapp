from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture(
    'IMG_2099.jpg', width=Inches(0.5)
)


name = input('What is your name? ')
speak('Hello ' + name + ' How are you today?')

speak('What is your phone number?')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About me')
about_me = input('Tell me about yourself!')
document.add_paragraph(about_me)

# work experience(s)
document.add_heading('Work experience')
p = document.add_paragraph()

company = input('Enter company ')
start_date = input('Start date ')
end_date = input('End date ')

p.add_run(company + ' ').bold = True
p.add_run(start_date + '-' + end_date + '\n').italic = True

experience_details = input(
    'Describe your work experience at ' + company + ' '
)
p.add_run(experience_details)

# more experiences
while True:
    more_experiences = input(
        'Do you have more experiences, Yes or No? ')
    if more_experiences.lower() == 'yes':
       p = document.add_paragraph()

       company = input('Enter company ')
       start_date = input('Start date ')
       end_date = input('End date ')

       p.add_run(company + ' ').bold = True
       p.add_run(start_date + '-' + end_date + '\n').italic = True

       experience_details = input(
       'Describe your work experience at ' + company
       )
       p.add_run(experience_details)
    else:
        break

# skills 
document.add_heading('Skills')
skills = input('Enter skills ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True: 
    more_skills = input('Do you have more skills, Yes or No? ')
    if more_skills.lower() == 'yes':
        skills = input('Enter extra skills ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Z.B institute."


document.save('CV.docx')